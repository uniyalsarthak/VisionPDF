import hashlib
from pathlib import Path
import fitz

from docling_core.types.doc import PictureItem, TextItem
from docling.datamodel.base_models import InputFormat
from docling.datamodel.pipeline_options import PdfPipelineOptions
from docling.document_converter import DocumentConverter, PdfFormatOption

from docx import Document
from docx.shared import Inches


# =========================================================
# GLOBAL DOCLING CONVERTER (INITIALIZED ONLY ONCE)
# =========================================================
pipeline_options = PdfPipelineOptions()
pipeline_options.images_scale = 2.5
pipeline_options.generate_picture_images = True
pipeline_options.generate_page_images = False
pipeline_options.do_ocr = False

CONVERTER = DocumentConverter(
    format_options={
        InputFormat.PDF: PdfFormatOption(pipeline_options=pipeline_options)
    }
)

# =========================================================
# FAST PAGE DETECTION (skip non-visual pages)
# =========================================================
def detect_pages_with_visuals(pdf_path):
    pdf = fitz.open(pdf_path)
    pages = []

    for i, page in enumerate(pdf):
        images = page.get_images(full=True)
        drawings = page.get_drawings()

        # skip decorative shapes
        if len(images) > 0 or len(drawings) > 10:
            pages.append(i)

    pdf.close()
    return pages


# =========================================================
# CREATE REDUCED PDF (ONLY IMPORTANT PAGES)
# =========================================================
def create_reduced_pdf(src_pdf, pages, out_pdf):
    src = fitz.open(src_pdf)
    dst = fitz.open()

    for p in pages:
        dst.insert_pdf(src, from_page=p, to_page=p)

    dst.save(out_pdf)
    src.close()
    dst.close()


# =========================================================
# CAPTION DETECTION HEURISTIC
# =========================================================
import re

def looks_like_caption(text):
    t = text.lower().strip()
    return re.match(r"^(fig|figure)\s*\.?\s*\d+", t) is not None


# =========================================================
# MAIN EXTRACTION FUNCTION (FAST)
# =========================================================
def extract_figures_with_captions(pdf_path):

    pdf_bytes = pdf_path.read_bytes()
    pdf_hash = hashlib.md5(pdf_bytes).hexdigest()

    out_dir = Path("extracted_images") / pdf_hash
    out_dir.mkdir(parents=True, exist_ok=True)

    # detect only pages that contain visuals
    pages = detect_pages_with_visuals(pdf_path)
    if not pages:
        return out_dir, []

    reduced_pdf = out_dir / "reduced.pdf"
    create_reduced_pdf(pdf_path, pages, reduced_pdf)

    # USE GLOBAL CONVERTER (VERY IMPORTANT)
    conv = CONVERTER.convert(str(reduced_pdf))
    items = list(conv.document.iterate_items())

    results = []
    idx = 0

    for i, (el, _) in enumerate(items):

        if isinstance(el, PictureItem):
            img = el.get_image(conv.document)
            if not img:
                continue

            caption = None

            for j in range(max(0, i - 3), min(len(items), i + 10)):
                nxt, _ = items[j]

                if isinstance(nxt, TextItem):
                    text = nxt.text.strip()

                    if looks_like_caption(text):
                        caption = text   # only single line
                        break

            idx += 1
            img_path = out_dir / f"figure_{idx}.png"
            img.save(img_path)

            results.append({
                "img_path": img_path,
                "caption": caption or "No caption"
            })

    return out_dir, results


# =========================================================
# WORD EXPORT
# =========================================================
def create_word_file(results, output_path, title="Extracted Figures"):

    doc = Document()
    doc.add_heading(title, level=1)

    for item in results:
        doc.add_picture(str(item["path"]), width=Inches(5))

        if item.get("caption"):
            doc.add_paragraph(item["caption"])
        else:
            doc.add_paragraph("Caption not detected.")

        doc.add_page_break()

    doc.save(output_path)
    return output_path
